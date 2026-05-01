"""
Parsers for:
  - DOCX files (python-docx, preserves tables)
  - Image files (JPEG, PNG, TIFF, BMP, WEBP → direct OCR)
"""

import logging
from io import BytesIO

from PIL import Image
from docx import Document
from docx.table import Table

from tender_platform.ingestion.ocr_engine import run_ocr
from tender_platform.models.document import FileFormat, PageResult

logger = logging.getLogger(__name__)


# ── DOCX ─────────────────────────────────────────────────────────────────────

def parse_docx(file_bytes: bytes, filename: str) -> tuple[list[PageResult], FileFormat, None]:
    """
    Parse a DOCX file.
    DOCX has no native page breaks accessible via python-docx,
    so we treat the whole document as a single logical page.
    Tables are extracted separately.
    Returns (pages, FileFormat.DOCX, None).
    """
    try:
        doc = Document(BytesIO(file_bytes))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        tables = []
        for table in doc.tables:
            grid = _extract_docx_table(table)
            if grid:
                tables.append(grid)
                # Append table as plain text too so the LLM sees it
                full_text += "\n" + _table_to_text(grid)

        pages = [PageResult(
            page_number=1,
            text=full_text.strip(),
            confidence=1.0,
            has_tables=bool(tables),
            tables=tables,
            is_ocr=False,
        )]
        return pages, FileFormat.DOCX, None

    except Exception as e:
        logger.error("DOCX parsing failed for '%s': %s", filename, e)
        raise


def _extract_docx_table(table: Table) -> list[list[str]]:
    grid = []
    for row in table.rows:
        grid.append([cell.text.strip() for cell in row.cells])
    return grid


def _table_to_text(grid: list[list[str]]) -> str:
    """Convert a 2D table into a pipe-separated markdown-ish string."""
    return "\n".join(" | ".join(row) for row in grid)


# ── Images ────────────────────────────────────────────────────────────────────

SUPPORTED_IMAGE_FORMATS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}


def parse_image(file_bytes: bytes, filename: str) -> tuple[list[PageResult], FileFormat, str]:
    """
    OCR a single image file.
    Returns (pages, FileFormat.IMAGE, ocr_engine_used).
    """
    try:
        img = Image.open(BytesIO(file_bytes))

        # Handle multi-frame images (TIFF stacks, animated GIFs)
        frames = _extract_frames(img)
        pages = []

        for i, frame in enumerate(frames):
            ocr_result = run_ocr(frame)
            page = PageResult(
                page_number=i + 1,
                text=ocr_result.text,
                confidence=ocr_result.confidence,
                is_ocr=True,
            )
            if ocr_result.confidence < 0.5:
                logger.warning(
                    "Very low OCR confidence (%.2f) on frame %d of image '%s'",
                    ocr_result.confidence, i + 1, filename,
                )
            pages.append(page)

        return pages, FileFormat.IMAGE, pages[0].is_ocr and ocr_result.engine or "none"

    except Exception as e:
        logger.error("Image parsing failed for '%s': %s", filename, e)
        raise


def _extract_frames(img: Image.Image) -> list[Image.Image]:
    """Extract all frames from a multi-page image (e.g., multi-page TIFF)."""
    frames = []
    try:
        i = 0
        while True:
            img.seek(i)
            frames.append(img.copy().convert("RGB"))
            i += 1
    except EOFError:
        pass
    return frames if frames else [img.convert("RGB")]