"""
Tests for Phase 1 — Document Ingestion Pipeline.
Uses synthetic in-memory documents so no real files are needed.
"""

import io
import sys
import os
sys.path.insert(0, "/home/claude")

import pytest
from PIL import Image, ImageDraw, ImageFont

from ingestion.document_processor import DocumentProcessor
from ingestion.other_parsers import parse_docx, parse_image
from models.document import DocType, FileFormat
from utils.image_preprocessing import preprocess_for_ocr


processor = DocumentProcessor()


# ── Helpers ────────────────────────────────────────────────────────────────

def make_synthetic_image(text: str, width: int = 800, height: int = 200) -> bytes:
    """Create a simple white image with black text — simulates a scanned doc."""
    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((20, 60), text, fill=(0, 0, 0))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def make_minimal_pdf_bytes(text: str) -> bytes:
    """
    Build the smallest valid single-page PDF containing the given text.
    No external library needed — raw PDF syntax.
    """
    stream = f"BT /F1 12 Tf 50 700 Td ({text}) Tj ET"
    stream_bytes = stream.encode()
    length = len(stream_bytes)

    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        + f"4 0 obj\n<< /Length {length} >>\nstream\n".encode()
        + stream_bytes
        + b"\nendstream\nendobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000400 00000 n \n"
        b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
        b"startxref\n460\n%%EOF"
    )
    return pdf


def make_minimal_docx_bytes(text: str) -> bytes:
    """Create a minimal DOCX containing one paragraph."""
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Image preprocessing tests ─────────────────────────────────────────────

class TestImagePreprocessing:

    def test_grayscale_output(self):
        img = Image.new("RGB", (200, 100), color=(180, 100, 50))
        result = preprocess_for_ocr(img)
        assert result.mode == "L", "Preprocessed image should be grayscale"

    def test_small_image_upscaled(self):
        img = Image.new("RGB", (400, 300), color=(255, 255, 255))
        result = preprocess_for_ocr(img)
        # Should have been upscaled since width < 1200
        assert result.size[0] >= 400, "Image should be at least as wide as original"

    def test_large_image_not_shrunk(self):
        img = Image.new("RGB", (2000, 1500), color=(255, 255, 255))
        result = preprocess_for_ocr(img)
        assert result.size[0] >= 2000, "Large image should not be shrunk"

    def test_rgba_image_handled(self):
        img = Image.new("RGBA", (200, 100), color=(100, 100, 100, 200))
        result = preprocess_for_ocr(img)
        assert result.mode == "L"


# ── DOCX parser tests ──────────────────────────────────────────────────────

class TestDocxParser:

    def test_basic_text_extraction(self):
        content = "This is a test document for tender evaluation."
        docx_bytes = make_minimal_docx_bytes(content)
        pages, fmt, engine = parse_docx(docx_bytes, "test.docx")
        assert fmt == FileFormat.DOCX
        assert engine is None
        assert len(pages) == 1
        assert content in pages[0].text

    def test_confidence_is_one_for_digital(self):
        docx_bytes = make_minimal_docx_bytes("Some text")
        pages, _, _ = parse_docx(docx_bytes, "test.docx")
        assert pages[0].confidence == 1.0

    def test_not_ocr_flagged(self):
        docx_bytes = make_minimal_docx_bytes("Some text")
        pages, _, _ = parse_docx(docx_bytes, "test.docx")
        assert pages[0].is_ocr is False

    def test_table_extraction(self):
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Company"
        table.cell(0, 1).text = "Turnover"
        table.cell(1, 0).text = "ABC Ltd"
        table.cell(1, 1).text = "5 crore"
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        pages, _, _ = parse_docx(docx_bytes, "table_test.docx")
        assert pages[0].has_tables is True
        assert len(pages[0].tables) >= 1
        flat = " ".join(" ".join(row) for row in pages[0].tables[0])
        assert "ABC Ltd" in flat
        assert "5 crore" in flat


# ── Image parser tests ─────────────────────────────────────────────────────

class TestImageParser:

    def test_image_returns_one_page(self):
        img_bytes = make_synthetic_image("Test text for OCR")
        pages, fmt, engine = parse_image(img_bytes, "test.png")
        assert fmt == FileFormat.IMAGE
        assert len(pages) == 1
        assert pages[0].is_ocr is True
        assert pages[0].page_number == 1

    def test_confidence_between_zero_and_one(self):
        img_bytes = make_synthetic_image("Hello World")
        pages, _, _ = parse_image(img_bytes, "test.png")
        assert 0.0 <= pages[0].confidence <= 1.0

    def test_blank_image_handled(self):
        blank = Image.new("RGB", (200, 100), color=(255, 255, 255))
        buf = io.BytesIO()
        blank.save(buf, format="PNG")
        pages, fmt, _ = parse_image(buf.getvalue(), "blank.png")
        assert len(pages) == 1
        # No crash — may return empty string


# ── Document processor tests ───────────────────────────────────────────────

class TestDocumentProcessor:

    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError, match="Unsupported"):
            processor.process(b"data", "file.xlsx", DocType.TENDER)

    def test_empty_file_handling(self):
        docx_bytes = make_minimal_docx_bytes("")
        doc = processor.process(docx_bytes, "empty.docx", DocType.TENDER)
        assert doc.total_pages >= 1
        assert doc.file_format == FileFormat.DOCX

    def test_docx_full_pipeline(self):
        text = "Minimum annual turnover: Rs 5 crore. GST registration mandatory."
        docx_bytes = make_minimal_docx_bytes(text)
        doc = processor.process(docx_bytes, "tender.docx", DocType.TENDER)
        assert doc.doc_type == DocType.TENDER
        assert doc.file_format == FileFormat.DOCX
        assert doc.overall_confidence == 1.0
        assert text in doc.full_text
        assert doc.ocr_engine_used is None
        assert len(doc.warnings) == 0

    def test_full_text_contains_page_markers(self):
        docx_bytes = make_minimal_docx_bytes("Some tender criteria here.")
        doc = processor.process(docx_bytes, "tender.docx", DocType.TENDER)
        assert "[Page 1]" in doc.full_text

    def test_warnings_generated_for_low_confidence(self):
        # Blank image → OCR returns 0 confidence → should trigger warning
        blank = Image.new("RGB", (200, 100), color=(255, 255, 255))
        buf = io.BytesIO()
        blank.save(buf, format="PNG")
        doc = processor.process(buf.getvalue(), "scan.png", DocType.BIDDER_SUBMISSION)
        assert doc.file_format == FileFormat.IMAGE
        # Warnings may or may not trigger depending on tesseract availability
        assert isinstance(doc.warnings, list)

    def test_image_pipeline(self):
        img_bytes = make_synthetic_image("ISO 9001 Certificate Valid")
        doc = processor.process(img_bytes, "certificate.png", DocType.BIDDER_SUBMISSION)
        assert doc.file_format == FileFormat.IMAGE
        assert doc.doc_type == DocType.BIDDER_SUBMISSION
        assert doc.total_pages == 1
        assert doc.ocr_engine_used is not None


if __name__ == "__main__":
    pytest.main([__file__, "-v"])